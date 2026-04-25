import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable


# ─── GERADOR DE WORD (.docx) ──────────────────────────────────────────────────

def gerar_docx(dados_hash: dict, dados_perito: dict, numero_caso: str, output_dir: str = ".") -> str:
    """
    Gera laudo pericial em formato Word (.docx).
    """

    doc = Document()

    # Configuração de margens
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2)

    # ── Cabeçalho ──────────────────────────────────────────────────────────────
    cabecalho = doc.add_paragraph()
    cabecalho.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cabecalho.add_run(dados_perito.get("instituicao", "INSTITUIÇÃO PERICIAL"))
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run("LAUDO PERICIAL DE ANÁLISE FORENSE COMPUTACIONAL")
    run2.bold = True
    run2.font.size = Pt(12)

    doc.add_paragraph()  # espaço

    # ── Linha divisória (tabela 1x1 com borda inferior) ───────────────────────
    tabela_linha = doc.add_table(rows=1, cols=1)
    tabela_linha.style = "Table Grid"
    cell = tabela_linha.cell(0, 0)
    cell.text = ""
    cell.paragraphs[0].add_run(f"Nº do Caso: {numero_caso}    |    Data: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

    doc.add_paragraph()

    # ── Dados do Perito ────────────────────────────────────────────────────────
    doc.add_heading("1. IDENTIFICAÇÃO DO PERITO RESPONSÁVEL", level=2)

    tabela_perito = doc.add_table(rows=3, cols=2)
    tabela_perito.style = "Table Grid"
    tabela_perito.alignment = WD_TABLE_ALIGNMENT.LEFT

    campos_perito = [
        ("Nome do Perito",   dados_perito.get("nome", "—")),
        ("Registro / CRP",   dados_perito.get("registro", "—")),
        ("Instituição",      dados_perito.get("instituicao", "—")),
    ]

    for i, (label, valor) in enumerate(campos_perito):
        row = tabela_perito.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = valor

    doc.add_paragraph()

    # ── Identificação da Evidência ─────────────────────────────────────────────
    doc.add_heading("2. IDENTIFICAÇÃO DA EVIDÊNCIA DIGITAL", level=2)

    tabela_evidencia = doc.add_table(rows=4, cols=2)
    tabela_evidencia.style = "Table Grid"

    campos_evidencia = [
        ("Nome do Arquivo",   dados_hash.get("arquivo", "—")),
        ("Caminho Completo",  dados_hash.get("caminho_completo", "—")),
        ("Tamanho",           dados_hash.get("tamanho_legivel", "—")),
        ("Data da Análise",   dados_hash.get("data_analise", "—")),
    ]

    for i, (label, valor) in enumerate(campos_evidencia):
        row = tabela_evidencia.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = valor

    doc.add_paragraph()

    # ── Valores de Hash ────────────────────────────────────────────────────────
    doc.add_heading("3. VALORES DE HASH (ASSINATURA DIGITAL)", level=2)

    tabela_hash = doc.add_table(rows=4, cols=2)
    tabela_hash.style = "Table Grid"

    # Cabeçalho da tabela
    hdr = tabela_hash.rows[0]
    hdr.cells[0].text = "Algoritmo"
    hdr.cells[1].text = "Valor do Hash"
    for cell in hdr.cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Cor de fundo do cabeçalho
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    for cell in hdr.cells:
        set_cell_bg(cell, "1A1A2E")

    hashes = [
        ("MD5",    dados_hash.get("md5", "—")),
        ("SHA-1",  dados_hash.get("sha1", "—")),
        ("SHA-256",dados_hash.get("sha256", "—")),
    ]

    for i, (algo, valor) in enumerate(hashes):
        row = tabela_hash.rows[i + 1]
        row.cells[0].text = algo
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = valor

    doc.add_paragraph()

    # ── Conclusão ──────────────────────────────────────────────────────────────
    doc.add_heading("4. CONCLUSÃO", level=2)
    conclusao = doc.add_paragraph()
    conclusao.add_run(
        "Os valores de hash acima representam a assinatura digital única do arquivo analisado "
        "e foram obtidos no momento da perícia. Qualquer alteração futura no arquivo resultará "
        "em valores de hash completamente diferentes, evidenciando adulteração da prova digital."
    )

    doc.add_paragraph()

    # ── Assinatura ─────────────────────────────────────────────────────────────
    doc.add_paragraph()
    assinatura = doc.add_paragraph()
    assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
    assinatura.add_run("_" * 50)

    ass2 = doc.add_paragraph()
    ass2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ass2.add_run(f"{dados_perito.get('nome', 'Perito Responsável')}\n")
    ass2.add_run(f"Registro: {dados_perito.get('registro', '—')}\n")
    ass2.add_run(f"{dados_perito.get('instituicao', '—')}")

    # ── Salvar ─────────────────────────────────────────────────────────────────
    os.makedirs(output_dir, exist_ok=True)
    nome_arquivo = f"laudo_{numero_caso.replace('/', '-')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    caminho_saida = os.path.join(output_dir, nome_arquivo)
    doc.save(caminho_saida)

    return caminho_saida


# ─── GERADOR DE PDF ───────────────────────────────────────────────────────────

def gerar_pdf(dados_hash: dict, dados_perito: dict, numero_caso: str, output_dir: str = ".") -> str:
    """
    Gera laudo pericial em formato PDF.
    """

    os.makedirs(output_dir, exist_ok=True)
    nome_arquivo = f"laudo_{numero_caso.replace('/', '-')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    caminho_saida = os.path.join(output_dir, nome_arquivo)

    doc = SimpleDocTemplate(
        caminho_saida,
        pagesize=A4,
        rightMargin=2*cm, leftMargin=3*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm
    )

    styles = getSampleStyleSheet()
    elementos = []

    AZUL_ESCURO = colors.HexColor("#1A1A2E")
    CINZA       = colors.HexColor("#555555")

    estilo_titulo = ParagraphStyle("titulo",
        parent=styles["Heading1"],
        fontSize=14, textColor=AZUL_ESCURO,
        alignment=1, spaceAfter=4)

    estilo_subtitulo = ParagraphStyle("subtitulo",
        parent=styles["Heading2"],
        fontSize=11, textColor=AZUL_ESCURO,
        spaceAfter=6, spaceBefore=12)

    estilo_normal = ParagraphStyle("normal",
        parent=styles["Normal"],
        fontSize=10, textColor=CINZA,
        spaceAfter=4, leading=14)

    # ── Cabeçalho ──────────────────────────────────────────────────────────────
    elementos.append(Paragraph(dados_perito.get("instituicao", "INSTITUIÇÃO PERICIAL"), estilo_titulo))
    elementos.append(Paragraph("LAUDO PERICIAL DE ANÁLISE FORENSE COMPUTACIONAL", estilo_titulo))
    elementos.append(HRFlowable(width="100%", thickness=1, color=AZUL_ESCURO))
    elementos.append(Spacer(1, 0.3*cm))
    elementos.append(Paragraph(
        f"<b>Nº do Caso:</b> {numero_caso} &nbsp;&nbsp;&nbsp; <b>Data:</b> {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}",
        estilo_normal))
    elementos.append(Spacer(1, 0.5*cm))

    # ── Dados do Perito ────────────────────────────────────────────────────────
    elementos.append(Paragraph("1. IDENTIFICAÇÃO DO PERITO RESPONSÁVEL", estilo_subtitulo))

    dados_tabela_perito = [
        ["Campo", "Informação"],
        ["Nome do Perito",  dados_perito.get("nome", "—")],
        ["Registro / CRP",  dados_perito.get("registro", "—")],
        ["Instituição",     dados_perito.get("instituicao", "—")],
    ]

    tabela_p = Table(dados_tabela_perito, colWidths=[5*cm, 11*cm])
    tabela_p.setStyle(TableStyle([
        ("BACKGROUND",  (0,0), (-1,0), AZUL_ESCURO),
        ("TEXTCOLOR",   (0,0), (-1,0), colors.white),
        ("FONTNAME",    (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE",    (0,0), (-1,-1), 9),
        ("GRID",        (0,0), (-1,-1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#F5F5F5")]),
        ("LEFTPADDING",  (0,0), (-1,-1), 8),
        ("RIGHTPADDING", (0,0), (-1,-1), 8),
        ("TOPPADDING",   (0,0), (-1,-1), 5),
        ("BOTTOMPADDING",(0,0), (-1,-1), 5),
    ]))
    elementos.append(tabela_p)
    elementos.append(Spacer(1, 0.5*cm))

    # ── Identificação da Evidência ─────────────────────────────────────────────
    elementos.append(Paragraph("2. IDENTIFICAÇÃO DA EVIDÊNCIA DIGITAL", estilo_subtitulo))

    dados_tabela_ev = [
        ["Campo", "Informação"],
        ["Nome do Arquivo",  dados_hash.get("arquivo", "—")],
        ["Caminho Completo", dados_hash.get("caminho_completo", "—")],
        ["Tamanho",          dados_hash.get("tamanho_legivel", "—")],
        ["Data da Análise",  dados_hash.get("data_analise", "—")],
    ]

    tabela_ev = Table(dados_tabela_ev, colWidths=[5*cm, 11*cm])
    tabela_ev.setStyle(TableStyle([
        ("BACKGROUND",  (0,0), (-1,0), AZUL_ESCURO),
        ("TEXTCOLOR",   (0,0), (-1,0), colors.white),
        ("FONTNAME",    (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE",    (0,0), (-1,-1), 9),
        ("GRID",        (0,0), (-1,-1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#F5F5F5")]),
        ("LEFTPADDING",  (0,0), (-1,-1), 8),
        ("RIGHTPADDING", (0,0), (-1,-1), 8),
        ("TOPPADDING",   (0,0), (-1,-1), 5),
        ("BOTTOMPADDING",(0,0), (-1,-1), 5),
        ("WORDWRAP",     (0,0), (-1,-1), True),
    ]))
    elementos.append(tabela_ev)
    elementos.append(Spacer(1, 0.5*cm))

    # ── Valores de Hash ────────────────────────────────────────────────────────
    elementos.append(Paragraph("3. VALORES DE HASH (ASSINATURA DIGITAL)", estilo_subtitulo))

    dados_tabela_hash = [
        ["Algoritmo", "Valor do Hash"],
        ["MD5",    dados_hash.get("md5", "—")],
        ["SHA-1",  dados_hash.get("sha1", "—")],
        ["SHA-256",dados_hash.get("sha256", "—")],
    ]

    tabela_h = Table(dados_tabela_hash, colWidths=[3*cm, 13*cm])
    tabela_h.setStyle(TableStyle([
        ("BACKGROUND",  (0,0), (-1,0), AZUL_ESCURO),
        ("TEXTCOLOR",   (0,0), (-1,0), colors.white),
        ("FONTNAME",    (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTNAME",    (0,1), (0,-1), "Helvetica-Bold"),
        ("FONTSIZE",    (0,0), (-1,-1), 8),
        ("GRID",        (0,0), (-1,-1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#F5F5F5")]),
        ("LEFTPADDING",  (0,0), (-1,-1), 8),
        ("RIGHTPADDING", (0,0), (-1,-1), 8),
        ("TOPPADDING",   (0,0), (-1,-1), 5),
        ("BOTTOMPADDING",(0,0), (-1,-1), 5),
    ]))
    elementos.append(tabela_h)
    elementos.append(Spacer(1, 0.5*cm))

    # ── Conclusão ──────────────────────────────────────────────────────────────
    elementos.append(Paragraph("4. CONCLUSÃO", estilo_subtitulo))
    elementos.append(Paragraph(
        "Os valores de hash acima representam a assinatura digital única do arquivo analisado "
        "e foram obtidos no momento da perícia. Qualquer alteração futura no arquivo resultará "
        "em valores de hash completamente diferentes, evidenciando adulteração da prova digital.",
        estilo_normal))
    elementos.append(Spacer(1, 1.5*cm))

    # ── Assinatura ─────────────────────────────────────────────────────────────
    elementos.append(HRFlowable(width="50%", thickness=0.5, color=AZUL_ESCURO, hAlign="CENTER"))
    elementos.append(Spacer(1, 0.2*cm))

    estilo_centro = ParagraphStyle("centro", parent=estilo_normal, alignment=1)
    elementos.append(Paragraph(f"<b>{dados_perito.get('nome', 'Perito Responsável')}</b>", estilo_centro))
    elementos.append(Paragraph(f"Registro: {dados_perito.get('registro', '—')}", estilo_centro))
    elementos.append(Paragraph(dados_perito.get("instituicao", "—"), estilo_centro))

    doc.build(elementos)
    return caminho_saida


# ─── Teste rápido ─────────────────────────────────────────────────────────────
if __name__ == "__main__":

    # Dados simulados para teste
    dados_hash = {
        "arquivo": "evidencia.txt",
        "caminho_completo": "/tmp/evidencia.txt",
        "tamanho_legivel": "24.00 B",
        "data_analise": "22/04/2026 18:54:48",
        "md5":    "5501428d2b3218e8384115e338fc0281",
        "sha1":   "7b7eb9021c9dd794b322e8b9444c9a864630d269",
        "sha256": "d680ff5ea44faa78c2254aa3ae7ab57f5379cae6c52e119f22314fcee86bd5db",
    }

    dados_perito = {
        "nome":        "Michele Silva",
        "registro":    "IPC-12345",
        "instituicao": "Instituto de Perícia Forense do Amazonas",
    }

    numero_caso = "2026/IPFA/001"

    print("\nGerando laudos...")

    caminho_docx = gerar_docx(dados_hash, dados_perito, numero_caso, output_dir="laudos")
    print(f"✅ Word gerado : {caminho_docx}")

    caminho_pdf = gerar_pdf(dados_hash, dados_perito, numero_caso, output_dir="laudos")
    print(f"✅ PDF gerado  : {caminho_pdf}")